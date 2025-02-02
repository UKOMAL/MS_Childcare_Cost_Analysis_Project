from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os

def create_doc():
    # Create a new Document
    doc = Document()
    
    # Read markdown content
    with open('shahid_dsc640_milestone4.md', 'r', encoding='utf-8') as file:
        content = file.read()

    # Process content
    sections = content.split('\n## ')
    
    # Add title and header info
    header_info = sections[0].split('\n')
    
    # Title
    title = header_info[0].replace('# ', '')
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata (author, course, etc.)
    for line in header_info[1:]:
        if line.strip():
            # Remove asterisks and clean up
            clean_line = line.replace('**', '').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_line)
            run.font.size = Pt(12)
    
    doc.add_paragraph()  # Add spacing

    # Define visualization paths and detailed descriptions
    visualizations = {
        'Initial Analysis': [
            {
                'path': '../milestone1/childcare_costs_map.png',
                'caption': 'Geographic Distribution of Childcare Costs Across the US',
                'description': 'This choropleth map visualizes the average childcare costs across different states, revealing significant regional variations. Darker shades indicate higher costs, with coastal states generally showing higher average expenses.'
            },
            {
                'path': '../milestone1/cost_distribution.png',
                'caption': 'Distribution of Childcare Costs by Age Group',
                'description': 'Box plots showing the distribution of childcare costs across different age groups. The visualization highlights that infant care consistently commands the highest costs, with median expenses ranging from $12,000 to $18,000 annually.'
            },
            {
                'path': '../milestone1/cost_trends.png',
                'caption': 'Childcare Cost Trends Over Time (2008-2018)',
                'description': 'Time series analysis showing the evolution of childcare costs over a decade. The trend lines demonstrate accelerating cost increases, particularly in urban areas, with a growing disparity between different regions.'
            },
            {
                'path': '../milestone1/urban_rural_comparison.png',
                'caption': 'Urban vs Rural Cost Comparison',
                'description': 'Comparative analysis of childcare costs between urban and rural areas. The visualization reveals unexpected patterns where some rural areas face disproportionate cost burdens relative to local income levels.'
            }
        ],
        'Advanced Analysis': [
            {
                'path': '../milestone3/output/network_viz.png',
                'caption': 'Network Analysis of Cost Relationships',
                'description': 'Interactive force-directed graph showing the complex relationships between various cost factors. Node sizes represent the strength of influence, while edge thicknesses indicate correlation strengths between variables.'
            },
            {
                'path': '../milestone3/output/bubble_viz.png',
                'caption': 'Cost Burden Analysis by State',
                'description': 'Bubble map overlay showing the relative cost burden across states. Bubble sizes represent the proportion of household income spent on childcare, while colors indicate absolute cost levels.'
            },
            {
                'path': '../milestone3/output/sankey_viz.png',
                'caption': 'Cost Flow Analysis Between Income Brackets',
                'description': 'Sankey diagram visualizing how childcare costs flow between different income brackets and regions. The width of the flows indicates the magnitude of cost burden on each demographic group.'
            },
            {
                'path': '../milestone3/output/3d_viz.png',
                'caption': '3D Analysis of Cost, Income, and Population',
                'description': 'Three-dimensional scatter plot exploring the relationship between childcare costs, median household income, and population density. The visualization reveals complex patterns that aren\'t visible in traditional 2D analyses.'
            },
            {
                'path': '../milestone3/output/static_costs.png',
                'caption': 'Comprehensive Cost Analysis Dashboard',
                'description': 'Static capture of our interactive dashboard showing multiple linked visualizations. The dashboard allows users to explore relationships between different variables and filter data based on various criteria.'
            },
            {
                'path': '../milestone3/output/infographic.png',
                'caption': 'Summary Infographic of Key Findings',
                'description': 'Visual summary of our key findings, including cost trends, regional variations, and demographic impacts. The infographic provides a quick overview of the most significant insights from our analysis.'
            }
        ]
    }

    # Process remaining sections
    for section in sections[1:]:
        if section.strip():
            # Split into title and content
            parts = section.split('\n', 1)
            if len(parts) == 2:
                title, content = parts
                
                # Add section heading
                doc.add_heading(title, 1)
                
                # Process subsections
                subsections = content.split('\n### ')
                
                # Add main section content
                main_content = subsections[0].strip()
                for paragraph in main_content.split('\n'):
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Add visualizations based on section
                if "Project Evolution" in title:
                    doc.add_heading('Initial Data Analysis', 2)
                    doc.add_paragraph("Our analysis began with exploring fundamental patterns in childcare costs across the United States. These visualizations established our baseline understanding and guided our subsequent investigations:")
                    
                    for viz in visualizations['Initial Analysis']:
                        if os.path.exists(viz['path']):
                            # Add visualization
                            doc.add_picture(viz['path'], width=Inches(6))
                            
                            # Add caption
                            caption = doc.add_paragraph()
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_run = caption.add_run(viz['caption'])
                            caption_run.bold = True
                            
                            # Add description
                            desc = doc.add_paragraph()
                            desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            desc.add_run(viz['description'])
                            
                            doc.add_paragraph()  # Spacing
                
                if "Technical Implementation" in title:
                    doc.add_heading('Interactive Dashboard Development', 2)
                    doc.add_paragraph("We developed a comprehensive interactive dashboard using Python, Plotly, and Dash to enable dynamic exploration of the data. The dashboard includes the following key features:")
                    
                    # Add dashboard features
                    features = [
                        "Interactive Geographic Analysis: Dynamic filtering and zooming capabilities for state-level exploration",
                        "Network Analysis Visualization: Real-time updates of relationship strengths and variable correlations",
                        "Dynamic Cost Flow Analysis: Interactive Sankey diagram showing cost distribution patterns",
                        "Multi-dimensional Analysis: 3D visualization with rotation and filtering capabilities",
                        "Comprehensive Filtering: Custom range selectors for costs, income levels, and geographic regions"
                    ]
                    for feature in features:
                        p = doc.add_paragraph(feature, style='List Bullet')
                    
                    # Add information about hosting options
                    doc.add_heading('Dashboard Accessibility', 2)
                    hosting_info = doc.add_paragraph()
                    hosting_info.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    hosting_info.add_run("The interactive dashboard can be accessed through multiple methods:\n\n")
                    
                    hosting_options = [
                        "Local Installation: Full interactive experience with complete data exploration capabilities",
                        "GitHub Repository: Access to source code and setup instructions",
                        "Video Demo: Walkthrough of key features and functionality (available upon request)",
                        "Static Captures: High-resolution images of key visualizations and insights"
                    ]
                    for option in hosting_options:
                        p = doc.add_paragraph(option, style='List Bullet')
                    
                    # Add advanced visualizations
                    doc.add_heading('Advanced Visualization Gallery', 2)
                    for viz in visualizations['Advanced Analysis']:
                        if os.path.exists(viz['path']):
                            # Add visualization
                            doc.add_picture(viz['path'], width=Inches(6))
                            
                            # Add caption
                            caption = doc.add_paragraph()
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_run = caption.add_run(viz['caption'])
                            caption_run.bold = True
                            
                            # Add description
                            desc = doc.add_paragraph()
                            desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            desc.add_run(viz['description'])
                            
                            doc.add_paragraph()  # Spacing
                
                # Process subsections
                for subsection in subsections[1:]:
                    sub_parts = subsection.split('\n', 1)
                    if len(sub_parts) == 2:
                        sub_title, sub_content = sub_parts
                        doc.add_heading(sub_title, 2)
                        
                        for paragraph in sub_content.strip().split('\n'):
                            if paragraph.strip():
                                p = doc.add_paragraph(paragraph.strip())
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save the document
    doc.save('shahid_dsc640_milestone4.docx')

if __name__ == "__main__":
    create_doc() 